"""Build a complete Computers & Electrical Engineering submission folder."""

from __future__ import annotations

from pathlib import Path
import json
import re
import shutil
import subprocess
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DESKTOP = Path.home() / "Desktop"
BASE_NAME = "Computers_and_Electrical_Engineering_Round2Revision_Riyang_Luo_2026-07-15"
DEST = DESKTOP / BASE_NAME
counter = 2
while DEST.exists():
    DEST = DESKTOP / f"{BASE_NAME}_{counter}"
    counter += 1

TITLE = "Support-gated selective recovery under controlled multi-sensor corruption"
AUTHOR = "Riyang Luo"
AFFILIATION = (
    "School of Mechanical and Power Engineering, Nanjing Tech University, "
    "30 South Puzhu Road, Nanjing 211816, Jiangsu Province, China"
)
EMAIL = "lry@njtech.edu.cn"
REPO = "https://github.com/Corey9999/recovery-supervision-calibrated-fallback"
REPO_DIR = DESKTOP / "recovery-supervision-calibrated-fallback"
RELEASE_TAG = "v1.5.0-round2-support-audit"
RELEASE_COMMIT = subprocess.check_output(
    ["git", "rev-list", "-n", "1", RELEASE_TAG], cwd=REPO_DIR, text=True
).strip()
COMPUTE_COMMIT = "59be82fc18aaf48b819c3c71a9482425f8baac45"

HIGHLIGHTS = [
    "A retrospectively audited support gate prevents 96.2% of harmful transfers.",
    "Only five of ten fitted endpoint pairs satisfy deployment support.",
    "All-row routing removes the correctness-disagreement selection restriction.",
    "Utility benefits vanish when one harm costs five or more corrections.",
    "Severity, sensor-group, mechanism and batch transport are audited.",
]
assert all(len(item) <= 85 for item in HIGHLIGHTS), [(len(item), item) for item in HIGHLIGHTS]

ABSTRACT = (
    "Recovering from sensor corruption requires choosing between a stable base classifier and a recovery endpoint that can both correct and damage decisions. "
    "We formulate this choice as controlled-corruption selective recovery. Endpoints and routing thresholds were fitted before evaluation on fixed future batches. "
    "The strict test set contains 1,436 observations per mechanism for which a controlled fault reached an available sensor group. "
    "The frozen conditional router prevented 94.7% of recovery-induced negative transfers while retaining 9.7% of available corrections. "
    "Only five of ten fitted endpoint pairs passed a deployment-support gate proposed after the frozen experiment, evaluated retrospectively and intended for prospective use. "
    "Applying that gate retrospectively reverted unsupported pairs to the base endpoint and yielded 96.2% prevention, 6.6% retention and mean macro-AUROC 0.8132, compared with 0.8115 for the base endpoint. "
    "Under equal correction and harm values, calibration-selected all-row and two-stage analytical controls yielded mean test utilities of 107.9 and 110.6 units per 10,000 observations; unconditional recovery yielded 132.1. "
    "On this fixed test set and with no compute penalty, the two-stage policy had the highest mean utility among the tested policies at a two-to-one harm ratio, but its 95% fitted-pair interval crossed zero. "
    "At ratios of five or ten every learned policy had negative mean test utility and the base endpoint dominated. "
    "Severity, sensor-group, mechanism and batch transport changed retention and utility materially. "
    "The evidence supports a reproducible risk-control protocol for controlled corruption, not natural-failure recovery or device-independent field performance."
)

KEYWORDS = "multi-sensor fusion; sensor corruption; fault-tolerant classification; selective recovery; negative transfer; cross-fitting"


def configure(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)


def heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)


def save_docx(path: Path, title: str, blocks: list[tuple[str | None, str | list[str]]]) -> None:
    doc = Document()
    configure(doc)
    heading(doc, title)
    for label, content in blocks:
        if label:
            run = doc.add_paragraph().add_run(label)
            run.bold = True
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(content)
    doc.core_properties.title = title
    doc.core_properties.author = AUTHOR
    doc.save(path)


def copy_source_tree(source_dir: Path) -> None:
    source_dir.mkdir(parents=True, exist_ok=True)
    for name in ("manuscript_cee.tex", "supplementary_cee.tex", "references.bib"):
        shutil.copy2(ROOT / name, source_dir / name)
    tables = source_dir / "tables"
    tables.mkdir()
    for path in (ROOT / "tables").glob("*.tex"):
        shutil.copy2(path, tables / path.name)

    tex_text = (ROOT / "manuscript_cee.tex").read_text(encoding="utf-8") + "\n" + (ROOT / "supplementary_cee.tex").read_text(encoding="utf-8")
    figure_names = sorted(set(re.findall(r"\\includegraphics(?:\[[^]]*\])?\{([^}]+)\}", tex_text)))
    figures = source_dir / "figures"
    figures.mkdir()
    for figure in figure_names:
        source = ROOT / "figures" / figure
        if source.suffix:
            candidates = [source]
        else:
            candidates = [source.with_suffix(ext) for ext in (".pdf", ".png")]
        selected = next((path for path in candidates if path.exists()), None)
        if selected is None:
            raise FileNotFoundError(f"Figure source not found: {figure}")
        shutil.copy2(selected, figures / selected.name)


def write_text(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def main() -> None:
    DEST.mkdir(parents=True)
    figures_dir = DEST / "Figures"
    figures_dir.mkdir()
    source_dir = DEST / "Source_Files"
    code_dir = DEST / "Reproducibility_Code"
    code_dir.mkdir()
    internal_dir = DEST / "Internal_QA_Not_For_Upload"
    internal_dir.mkdir()

    shutil.copy2(ROOT / "manuscript_cee_editable.docx", DEST / "02_Manuscript_CEE.docx")
    shutil.copy2(ROOT / "manuscript_cee.pdf", DEST / "02_Manuscript_CEE.pdf")
    shutil.copy2(ROOT / "supplementary_cee_editable.docx", DEST / "09_Supplementary_Information.docx")
    shutil.copy2(ROOT / "supplementary_cee.pdf", DEST / "09_Supplementary_Information.pdf")

    save_docx(
        DEST / "01_Title_Page.docx",
        "Title page",
        [
            ("Article title", TITLE),
            ("Article type", "Original research article"),
            ("Author", AUTHOR),
            ("Affiliation", AFFILIATION),
            ("Corresponding author", f"{AUTHOR}; Email: {EMAIL}"),
            ("Running title", "Support-gated recovery under controlled corruption"),
            ("Funding", "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors."),
            ("Declaration of interests", "The author has nothing to declare."),
        ],
    )
    save_docx(DEST / "03_Highlights.docx", "Highlights", [(None, HIGHLIGHTS)])
    save_docx(
        DEST / "04_Cover_Letter.docx",
        "Cover letter",
        [
            (None, "15 July 2026"),
            (None, "Editor-in-Chief\nComputers & Electrical Engineering"),
            (None, "Dear Editor-in-Chief,"),
            (
                None,
                f"Please consider the original research article entitled \"{TITLE}\" for publication in Computers & Electrical Engineering. The manuscript develops and audits a computational risk-control method for fault-tolerant multi-sensor classification, matching the journal's scope in signal processing, computational methods and dependable electrical and information systems.",
            ),
            (
                None,
                "The principal contribution is a two-pass selective-recovery decision framework that distinguishes a conditional disagreement score from an all-row deployment estimand. Always-base and always-recovery policies, complete prevention-retention frontiers and calibration-matched simple controls establish the decision context. A support gate proposed after the frozen experiment is audited retrospectively over 27 threshold configurations; the reference configuration passes five of ten endpoint pairs and reverts unsupported pairs to the base endpoint, yielding 96.2% negative-transfer prevention and 6.6% correction retention. All-row and two-stage models remove the selected-disagreement restriction but remain analytical controls without model-specific deployment gates. Fitted-pair utility intervals and calibration-to-test changes qualify the cost-sensitive results. Severity, affected-group, leave-one-mechanism-out and complete acquisition-batch analyses expose transport limits. Claims remain limited to controlled available-group corruption; natural-failure and device-independent claims are not made.",
            ),
            (
                None,
                f"The code, source data and figure/table scripts are available at {REPO}. The frozen computational snapshot is commit {COMPUTE_COMMIT}; release tag {RELEASE_TAG} targets {RELEASE_COMMIT}. A software DOI has not yet been assigned. All 62 cited references were checked against DOI or official publication records.",
            ),
            (
                None,
                "This manuscript is original, has not been published previously and is not under consideration elsewhere. I am the sole author and approve the submitted version. The work received no specific funding, and I have no competing interests.",
            ),
            (None, f"Sincerely,\n{AUTHOR}\n{AFFILIATION}\n{EMAIL}"),
        ],
    )
    save_docx(
        DEST / "05_Declaration_of_Interest.docx",
        "Declaration of interests",
        [(None, "The author declares that there are no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")],
    )
    save_docx(
        DEST / "06_CRediT_Author_Statement.docx",
        "CRediT author statement",
        [(None, "Riyang Luo: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Visualization, Writing - original draft, Writing - review and editing, and Project administration.")],
    )
    save_docx(
        DEST / "07_Funding_Statement.docx",
        "Funding statement",
        [(None, "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.")],
    )
    save_docx(
        DEST / "08_Data_and_Code_Availability.docx",
        "Data and code availability",
        [
            (None, "Gas Sensor Array Drift Dataset: https://doi.org/10.24432/C5RP6W"),
            (None, "Condition Monitoring of Hydraulic Systems: https://doi.org/10.24432/C5CW21"),
            (None, "Operational AHU dataset: https://doi.org/10.6084/m9.figshare.27147678.v3"),
            (None, f"Code repository: {REPO}"),
            (None, f"Frozen computational snapshot: {COMPUTE_COMMIT}"),
            (None, f"Versioned release: {RELEASE_TAG} (target commit {RELEASE_COMMIT})"),
            (None, "CEE revision directory: cee_revision/"),
            (None, "A permanent software DOI has not been assigned."),
        ],
    )
    captions = [
        "Figure 1. Bounded responsive fusion and recovery supervision. Panel a shows the bounded group-fusion path; panel b shows shared clean/faulted parameters and recovery losses.",
        "Figure 2. Integrated experimental protocol and controlled fault design. The chemical-array temporal split and independent hydraulic guarded-block protocol are shown with their intervention boundaries.",
        "Figure 3. Selective recovery as an explicit decision policy. The panels report complete prevention-retention curves, fixed policy context, utility across harm ratios, and the all-row out-of-fold confusion matrix.",
        "Figure 4. Selector transport and acquisition-batch robustness. The panels report severity transport, affected-group transport, observation/equal/worst-batch views, and leave-one-mechanism-out utility.",
        "Figure 5. Selector stability and final probability quality. The panels show repeated grouped-validation dispersion, coefficient sign stability, reliability curves and endpoint-agreement-stratified NLL/ECE.",
    ]
    save_docx(DEST / "11_Figure_Captions.docx", "Figure captions", [(None, captions)])

    figure_map = {
        "Figure_1.pdf": "figure1_method_framework.pdf",
        "Figure_2.pdf": "figure2_experimental_protocol.pdf",
        "Figure_3.pdf": "figure3_selective_recovery_decision.pdf",
        "Figure_4.pdf": "figure4_shift_batch_robustness.pdf",
        "Figure_5.pdf": "figure_cee_q1_stability_calibration.pdf",
    }
    for target, source in figure_map.items():
        shutil.copy2(ROOT / "figures" / source, figures_dir / target)
        tiff_source = (ROOT / "figures" / source).with_suffix(".tiff")
        if tiff_source.exists():
            shutil.copy2(tiff_source, figures_dir / Path(target).with_suffix(".tiff").name)
    shutil.copy2(ROOT / "figures" / "graphical_abstract_cee.pdf", DEST / "10_Graphical_Abstract_OPTIONAL.pdf")
    shutil.copy2(ROOT / "figures" / "graphical_abstract_cee.tiff", DEST / "10_Graphical_Abstract_OPTIONAL.tiff")

    copy_source_tree(source_dir)
    for name in (
        "run_cee_selector_validation.py",
        "run_cee_lite_routing_validation.py",
        "analyse_cee_results.py",
        "analyse_cee_q1_scores.py",
        "analyse_cee_lite_results.py",
        "analyse_cee_fault_plausibility.py",
        "run_cee_strict88_additional_audits.py",
        "analyse_cee_strict88_revision.py",
        "analyse_cee_major_router_audits.py",
        "run_cee_major_severity_transport.py",
        "analyse_cee_major_shift_audits.py",
        "analyse_cee_round2_revision.py",
        "make_cee_major_revision_figures.py",
        "make_cee_q1_figures.py",
        "make_cee_validation_figure.py",
        "make_cee_graphical_abstract.py",
        "STRICT88_FIGURE_CONTRACT.md",
        "STRICT88_REVISION_AUDIT.md",
        "MAJOR_ROUTER_FIGURE_CONTRACT.md",
        "MAJOR_REVISION_AUDIT.md",
        "ROUND2_REVISION_AUDIT.md",
        "GRAPHICAL_ABSTRACT_CONTRACT.md",
        "requirements-lock.txt",
        "environment.yml",
        "cee_cf10_r2_lite_config.json",
    ):
        shutil.copy2(ROOT / name, code_dir / name)
    code_tables = code_dir / "tables"
    code_tables.mkdir()
    for path in sorted((ROOT / "tables").glob("round2_*.tex")):
        shutil.copy2(path, code_tables / path.name)
    output_data = code_dir / "source_data"
    output_data.mkdir()
    for path in sorted((ROOT / "source_data").glob("q1_*")):
        if path.is_file():
            shutil.copy2(path, output_data / path.name)
    for path in sorted((ROOT / "source_data").glob("strict88_*")):
        if path.is_file():
            shutil.copy2(path, output_data / path.name)
    for path in sorted((ROOT / "source_data").glob("major_*")):
        if path.is_file():
            shutil.copy2(path, output_data / path.name)
    for path in sorted((ROOT / "source_data").glob("round2_*")):
        if path.is_file():
            shutil.copy2(path, output_data / path.name)
    for name in (
        "cee_cf10_design.json",
        "cee_cf10_summary.json",
        "cee_cf10_recovery_eligibility.csv",
        "cee_cf10_selector_calibration.csv.gz",
        "cee_cf10_strict_predictions.csv.gz",
        "cee_cf10_stream_safety.csv",
    ):
        shutil.copy2(ROOT / "source_data" / name, output_data / name)

    archive = DEST / "Source_Files.zip"
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zf:
        for path in source_dir.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(source_dir))

    word_count = len(re.findall(r"\b[\w-]+\b", ABSTRACT))
    write_text(
        DEST / "00_READ_ME_FIRST.txt",
        f"""
Target journal: Computers & Electrical Engineering
Article type: Original research article
Author: {AUTHOR} (sole author)

Recommended upload order:
1. 02_Manuscript_CEE.docx (editable main manuscript; author details retained because CEE uses single-anonymized review)
2. 01_Title_Page.docx
3. 03_Highlights.docx (required; five bullets, each <= 85 characters)
4. 04_Cover_Letter.docx
5. 05_Declaration_of_Interest.docx (required separate Word declaration)
6. 09_Supplementary_Information.pdf and/or editable DOCX
7. Figures/Figure_1.pdf through Figure_5.pdf (TIFF companions are also supplied)
8. 10_Graphical_Abstract_OPTIONAL.pdf or TIFF (encouraged, not required)
9. Source_Files.zip if the system requests LaTeX source files

The manuscript abstract contains {word_count} words (CEE limit: 250).
The graphical abstract is 1328 x 531 pixels in PNG layout and supplied as PDF/TIFF.
The frozen computational snapshot is {COMPUTE_COMMIT}.
Release {RELEASE_TAG} points to {RELEASE_COMMIT}.

Do not upload the folder Internal_QA_Not_For_Upload.
""",
    )
    write_text(
        DEST / "Code_Repository.txt",
        f"Repository: {REPO}\nFrozen computational snapshot: {COMPUTE_COMMIT}\nRelease: {RELEASE_TAG}\nRelease target commit: {RELEASE_COMMIT}\nRevision directory: cee_revision/\nSoftware DOI: not assigned\n",
    )
    write_text(DEST / "Abstract_and_Keywords.txt", f"Abstract\n\n{ABSTRACT}\n\nKeywords\n\n{KEYWORDS}\n")

    checklist = f"""
# CEE submission checklist

- [x] Editable single-column Word manuscript prepared.
- [x] Author details retained; CEE uses single-anonymized peer review.
- [x] Full postal affiliation and corresponding-author email supplied.
- [x] Abstract is {word_count} words (limit 250) and contains no references.
- [x] Six English keywords supplied (allowed range 1-7).
- [x] Five highlights supplied; every bullet is at most 85 characters.
- [x] Numeric references used; 62 of 62 cited records verified.
- [x] Tables are editable and placed near relevant text.
- [x] Five main figures supplied separately as editable vector PDFs and high-resolution TIFFs.
- [x] Optional graphical abstract supplied at the required 1328 x 531 proportion.
- [x] Separate Word declaration of interests supplied.
- [x] No-funding statement supplied.
- [x] CRediT statement supplied for the sole author.
- [x] Round-two computational snapshot and release tag verified on the public remote.
- [x] Supplementary information supplied in PDF and editable Word formats.
- [x] Main and supplementary PDFs compile without undefined citations/references.
- [x] Page-level visual QA found no overlapping arrows, cropped labels or out-of-frame text.
- [ ] Complete the Elsevier online declarations tool and select "I have nothing to declare".
- [ ] Enter author metadata in the submission portal and confirm spelling before final submission.
"""
    write_text(DEST / "Submission_Checklist.md", checklist)

    shutil.copy2(ROOT / "MAJOR_REVISION_AUDIT.md", internal_dir / "MAJOR_REVISION_AUDIT_INTERNAL.md")
    shutil.copy2(ROOT / "ROUND2_REVISION_AUDIT.md", internal_dir / "ROUND2_REVISION_AUDIT_INTERNAL.md")
    shutil.copy2(ROOT / "REFERENCE_VERIFICATION_62.md", internal_dir / "REFERENCE_VERIFICATION_62_INTERNAL.md")
    qa_summary = {
        "main_pdf_pages": 33,
        "supplement_pdf_pages": 50,
        "main_word_render_pages": 40,
        "supplement_word_render_pages": 65,
        "main_word_render_checked": True,
        "latex_undefined_references": 0,
        "latex_undefined_citations": 0,
        "references_verified": 62,
        "computational_snapshot": COMPUTE_COMMIT,
        "release_target_commit": RELEASE_COMMIT,
        "release_tag": RELEASE_TAG,
    }
    write_text(internal_dir / "QA_SUMMARY.json", json.dumps(qa_summary, indent=2))

    print(DEST)


if __name__ == "__main__":
    main()

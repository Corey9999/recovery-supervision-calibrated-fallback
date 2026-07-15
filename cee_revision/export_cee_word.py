"""Convert CEE HTML outputs to editable, single-column Word documents."""

from pathlib import Path
import re
import sys

ROOT = Path(__file__).resolve().parent
TOOLS = ROOT.parent / "paper_package_q1_upgrade" / ".word_tools"
sys.path.insert(0, str(TOOLS))
import pypandoc  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Mm, Pt  # noqa: E402


def convert(source: str, target: str) -> Path:
    destination = ROOT / target
    pypandoc.convert_file(
        str(ROOT / source),
        to="docx",
        format="html",
        outputfile=str(destination),
        extra_args=["--resource-path=" + str(ROOT), "--standalone"],
    )
    return destination


def normalized(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u00a0", " ")).strip()


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_caption(doc, prefix: str):
    for paragraph in doc.paragraphs:
        if normalized(paragraph.text).startswith(prefix):
            return paragraph
    raise RuntimeError(f"Caption {prefix!r} was not found")


def replace_converted_figure(doc, caption_prefix: str, image_name: str) -> None:
    caption = find_caption(doc, caption_prefix)
    node = caption._p.getprevious()
    removable = []
    found_image = False
    while node is not None:
        previous = node.getprevious()
        if node.tag == qn("w:tbl") and node.xpath('.//*[local-name()="blip"]'):
            removable.append(node)
            found_image = True
        elif node.tag == qn("w:p") and not "".join(node.itertext()).strip():
            removable.append(node)
        else:
            break
        node = previous
    if found_image:
        for element in removable:
            element.getparent().remove(element)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ROOT / "figures" / image_name), width=Inches(6.25))
    caption._p.addprevious(paragraph._p)


def set_repeat_header_and_no_split(table) -> None:
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)


def set_cell_width(cell, width_inches: float) -> None:
    """Apply a fixed cell width in both the python-docx and OOXML layers."""
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width.twips)))
    tc_width.set(qn("w:type"), "dxa")
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        margin = margins.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            margins.append(margin)
        margin.set(qn("w:w"), "25")
        margin.set(qn("w:type"), "dxa")


def format_wide_utility_tables(doc) -> None:
    """Keep the ten-column fitted-pair utility tables legible in portrait Word."""
    widths = (0.44, 0.76, 0.45, 1.00, 1.00, 0.57, 0.51, 0.44, 0.54, 0.49)
    header_labels = (
        "Ratio", "Policy", "Cal.\nU", "Test U\n[95% CI]", "Median\n[IQR]",
        "Test-cal.", "Positive\npairs", "Select\n(%)", "Prevent\n(%)", "Retain\n(%)",
    )
    for table in doc.tables:
        if not table.rows or len(table.rows[0].cells) != len(widths):
            continue
        headers = [normalized(cell.text) for cell in table.rows[0].cells]
        signature = " | ".join(headers)
        if not (headers[0].startswith("Ratio") and "Cal." in signature and "Prevent" in signature):
            continue
        table.autofit = False
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table._tbl.tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")
        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        for grid_column, width in zip(grid_columns, widths):
            grid_column.set(qn("w:w"), str(int(Inches(width).twips)))
        for cell, label in zip(table.rows[0].cells, header_labels):
            cell.text = label
        for row in table.rows:
            for index, (cell, width) in enumerate(zip(row.cells, widths)):
                set_cell_width(cell, width)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if index == 1 else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    for run in paragraph.runs:
                        run.font.size = Pt(5.5)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def move_table_captions(doc) -> None:
    for paragraph in list(doc.paragraphs):
        if not re.match(r"^Table\s*\d+:", normalized(paragraph.text)):
            continue
        p_pr = paragraph._p.get_or_add_pPr()
        p_pr.append(OxmlElement("w:keepNext"))
        p_pr.append(OxmlElement("w:keepLines"))
        node = paragraph._p.getprevious()
        while node is not None and node.tag == qn("w:p") and not "".join(node.itertext()).strip():
            node = node.getprevious()
        if node is not None and node.tag == qn("w:tbl"):
            node.addprevious(paragraph._p)


def clean_latex_cell(text: str) -> str:
    text = text.strip().replace(r"\%", "%").replace("$", "")
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    return text


def replace_main_utility_table(doc) -> None:
    """Replace TeX4ht's nested tabular with an editable fixed-grid Word table."""
    caption = find_caption(doc, "Table 9:")
    source = (ROOT / "tables" / "round2_utility_fitted_pair_main.tex").read_text(
        encoding="utf-8"
    )
    body = source.split(r"\midrule", 1)[1].split(r"\bottomrule", 1)[0]
    rows = []
    for line in body.splitlines():
        line = line.strip()
        if not line or "&" not in line:
            continue
        line = re.sub(r"\\\\\s*$", "", line)
        rows.append([clean_latex_cell(value) for value in line.split("&")])

    converted = caption._p.getprevious()
    while converted is not None and converted.tag == qn("w:p") and not "".join(converted.itertext()).strip():
        converted = converted.getprevious()
    if converted is None or converted.tag != qn("w:tbl"):
        converted = caption._p.getnext()
        while converted is not None and converted.tag == qn("w:p") and not "".join(converted.itertext()).strip():
            converted = converted.getnext()
    if converted is None or converted.tag != qn("w:tbl"):
        raise RuntimeError("Converted Table 9 tabular was not found after its caption")
    converted.getparent().remove(converted)

    table = doc.add_table(rows=1, cols=10)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "3")
        border.set(qn("w:color"), "B7B7B7")
        borders.append(border)
    table._tbl.tblPr.append(borders)
    headers = (
        "Ratio", "Policy", "Cal. U", "Test U [95% CI]", "Median [IQR]",
        "Test-cal.", "Positive", "Select (%)", "Prevent (%)", "Retain (%)",
    )
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    set_repeat_header_and_no_split(table)
    caption._p.addnext(table._tbl)


def style_document(doc, line_numbers: bool) -> None:
    for section in doc.sections:
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        if line_numbers:
            sect_pr = section._sectPr
            existing = sect_pr.find(qn("w:lnNumType"))
            if existing is not None:
                sect_pr.remove(existing)
            line_num = OxmlElement("w:lnNumType")
            line_num.set(qn("w:countBy"), "1")
            line_num.set(qn("w:restart"), "continuous")
            line_num.set(qn("w:distance"), "360")
            sect_pr.append(line_num)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    for style_name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(size)

    max_width = Inches(6.25)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.height = int(shape.height * ratio)
            shape.width = max_width
    for table in doc.tables:
        table.autofit = True
        set_repeat_header_and_no_split(table)
        columns = max((len(row.cells) for row in table.rows), default=1)
        font_size = Pt(7.5 if columns >= 7 else 8.5 if columns >= 5 else 9)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1
                    for run in paragraph.runs:
                        run.font.size = font_size
                        run.font.name = "Times New Roman"


def remove_unused_images(doc) -> None:
    used = {node.get(qn("r:embed")) for node in doc.element.xpath('.//*[local-name()="blip"]')}
    for rel_id, relationship in list(doc.part.rels.items()):
        if relationship.reltype == RT.IMAGE and rel_id not in used:
            doc.part.drop_rel(rel_id)


def prepare_main(path: Path) -> None:
    doc = Document(path)
    figures = [
        ("Figure 1:", "figure1_method_framework.png"),
        ("Figure 2:", "figure2_experimental_protocol.png"),
        ("Figure 3:", "figure3_selective_recovery_decision.png"),
        ("Figure 4:", "figure4_shift_batch_robustness.png"),
        ("Figure 5:", "figure_cee_q1_stability_calibration.png"),
    ]
    for caption, image in figures:
        replace_converted_figure(doc, caption, image)
    replace_main_utility_table(doc)
    move_table_captions(doc)
    style_document(doc, line_numbers=True)
    format_wide_utility_tables(doc)

    title = "Support-gated selective recovery under controlled multi-sensor corruption"
    duplicates = [p for p in doc.paragraphs[:5] if normalized(p.text) == title]
    for paragraph in duplicates[1:]:
        delete_paragraph(paragraph)
    paragraphs = list(doc.paragraphs)
    title_index = next(i for i, p in enumerate(paragraphs) if normalized(p.text) == title)
    abstract_index = next(i for i, p in enumerate(paragraphs) if normalized(p.text) == "Abstract")
    first = paragraphs[title_index]
    # TeX4ht emits the author block and a synthetic line-number sequence as
    # ordinary Word paragraphs.  Remove everything between the retained title
    # and Abstract, then insert one clean author block for submission.
    for paragraph in paragraphs[title_index + 1 : abstract_index]:
        delete_paragraph(paragraph)
    first.style = doc.styles["Title"]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run(
        "Riyang Luo*\n"
        "School of Mechanical and Power Engineering, Nanjing Tech University, "
        "30 South Puzhu Road, Nanjing 211816, Jiangsu Province, China\n"
        "*Corresponding author: lry@njtech.edu.cn"
    )
    first._p.addnext(author._p)
    remove_unused_images(doc)
    doc.core_properties.title = title
    doc.core_properties.author = "Riyang Luo"
    doc.save(path)


def prepare_supplement(path: Path) -> None:
    doc = Document(path)
    move_table_captions(doc)
    style_document(doc, line_numbers=False)
    format_wide_utility_tables(doc)
    title = (
        "Supplementary information Support-gated selective recovery under "
        "controlled multi-sensor corruption"
    )
    duplicates = [p for p in doc.paragraphs[:6] if normalized(p.text) == title]
    for paragraph in duplicates[1:]:
        delete_paragraph(paragraph)
    if duplicates:
        duplicates[0].text = (
            "Supplementary information\n"
            "Support-gated selective recovery under controlled multi-sensor corruption"
        )
        duplicates[0].style = doc.styles["Title"]
        duplicates[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = next((p for p in doc.paragraphs[:6] if normalized(p.text) == "Riyang Luo"), None)
    if author is not None:
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.core_properties.title = title
    doc.core_properties.author = "Riyang Luo"
    doc.save(path)


if __name__ == "__main__":
    main_doc = convert("manuscript_cee.html", "manuscript_cee_editable.docx")
    supp_doc = convert("supplementary_cee.html", "supplementary_cee_editable.docx")
    prepare_main(main_doc)
    prepare_supplement(supp_doc)
    for path in (main_doc, supp_doc):
        document = Document(path)
        print(f"{path.name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables, {len(document.inline_shapes)} figures")
